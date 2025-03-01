import os
from openai import OpenAI
from mistralai import Mistral
from docx import Document
from docx.shared import Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image
import io
from dotenv import load_dotenv
import re
import ast

load_dotenv()

openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
mistral_client = Mistral(api_key=os.getenv("MISTRAL_API_KEY"))


def generate_report(
    project_name,
    sections,
    descriptions,
    total_word_limit,
    image_mapping,
    uploaded_images,
    font,
    color,
    include_toc,
    citations,
):
    doc = Document()
    ensure_toc_styles(doc)
    set_document_style(doc, font, color)

    if include_toc:
        add_table_of_contents(doc, sections)

    word_limit_per_section = total_word_limit // len(sections)

    for section in sections:
        initial_content = generate_initial_content(
            project_name, section, descriptions[section], word_limit_per_section
        )
        enhanced_content = enhance_content(initial_content, word_limit_per_section)

        doc.add_heading(section, level=1)
        paragraph = doc.add_paragraph(enhanced_content)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_images_to_section(
            doc, section, image_mapping, uploaded_images, project_name
        )

    if citations:
        add_citations(doc, citations)

    return save_document(doc)


def generate_initial_content(project_name, section, description, word_limit):
    response = openai_client.chat.completions.create(
        model="gpt-4",
        messages=[
            {
                "role": "system",
                "content": "You are a report writer generating content for a specific section.",
            },
            {
                "role": "user",
                "content": f"Generate a {word_limit}-word {section} for a report on the project '{project_name}' based on this description: {description}",
            },
        ],
    )
    return response.choices[0].message.content


def enhance_content(initial_content, target_word_count):
    current_word_count = len(initial_content.split())
    if current_word_count >= target_word_count:
        return initial_content

    response = mistral_client.chat.complete(
        model="mistral-large-latest",
        messages=[
            {
                "role": "system",
                "content": "You are an AI assistant tasked with expanding and enhancing text to meet a specific word count.",
            },
            {
                "role": "user",
                "content": f"Expand the following text to approximately {target_word_count} words while maintaining coherence and adding relevant details:\n\n{initial_content}",
            },
        ],
    )
    return response.choices[0].message.content


def ensure_toc_styles(doc):
    for i in range(1, 4):
        toc_style_name = f"TOC {i}"
        if toc_style_name not in doc.styles:
            doc.styles.add_style(toc_style_name, WD_STYLE_TYPE.PARAGRAPH)


def set_document_style(doc, font, color):
    style = doc.styles["Normal"]
    style.font.name = font
    style.font.color.rgb = RGBColor.from_string(color[1:])


def add_table_of_contents(doc, sections):
    doc.add_heading("Table of Contents", level=1)
    for section in sections:
        doc.add_paragraph(section, style="TOC 1")
    doc.add_page_break()


def add_images_to_section(doc, section, image_mapping, uploaded_images, project_name):
    for image_name, mapped_section in image_mapping.items():
        if mapped_section == section:
            image = next(
                (img for img in uploaded_images if img.name == image_name), None
            )
            if image:
                img_stream = io.BytesIO(image.getvalue())
                doc.add_picture(img_stream, width=Inches(6))
                caption = generate_image_caption(project_name, section, image_name)
                doc.add_paragraph(caption, style="Caption")


def generate_image_caption(project_name, section, image_name):
    response = mistral_client.chat.complete(
        model="mistral-large-latest",
        messages=[
            {
                "role": "system",
                "content": "You are an AI assistant that generates image captions.",
            },
            {
                "role": "user",
                "content": f"Generate a brief, informative caption for an image named '{image_name}' in the {section} section of a report about the project '{project_name}'.",
            },
        ],
    )
    return response.choices[0].message.content


def add_citations(doc, citations):
    doc.add_heading("References", level=1)
    for citation in citations.split("\n"):
        doc.add_paragraph(citation.strip(), style="List Bullet")


def save_document(doc):
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream


def analyze_code_blocks(code_blocks):
    """
    Analyze a list of code blocks to identify functions, classes, and their relationships.
    Returns structured data that can be used to generate architecture diagrams.
    """
    analysis_results = []

    for i, block in enumerate(code_blocks):
        block_analysis = {
            "id": f"block_{i}",
            "type": "unknown",
            "name": f"Block {i + 1}",
            "functions": [],
            "classes": [],
            "imports": [],
        }

        # Extract functions and classes using regex
        function_matches = re.findall(r"def\s+([a-zA-Z_][a-zA-Z0-9_]*)\s*\(", block)
        class_matches = re.findall(r"class\s+([a-zA-Z_][a-zA-Z0-9_]*)\s*[:\(]", block)
        import_matches = re.findall(
            r"(?:import|from)\s+([a-zA-Z_][a-zA-Z0-9_\.]*)", block
        )

        if function_matches:
            block_analysis["functions"] = function_matches
            block_analysis["type"] = "function_container"

        if class_matches:
            block_analysis["classes"] = class_matches
            block_analysis["type"] = "class_container"

        if import_matches:
            block_analysis["imports"] = import_matches

        # Try to give it a meaningful name
        if function_matches:
            block_analysis["name"] = function_matches[0]
        elif class_matches:
            block_analysis["name"] = class_matches[0]
        else:
            # Use first non-empty line as name
            lines = [line for line in block.split("\n") if line.strip()]
            if lines:
                block_analysis["name"] = lines[0][:30] + (
                    "..." if len(lines[0]) > 30 else ""
                )

        analysis_results.append(block_analysis)

    return analysis_results


def generate_advanced_mermaid(code_blocks):
    """
    Generate a more sophisticated Mermaid diagram from code blocks,
    identifying functions, classes, and their relationships.
    """
    analysis = analyze_code_blocks(code_blocks)

    mermaid_code = "flowchart TD\n"

    # Add nodes
    for block in analysis:
        node_id = block["id"]
        node_name = block["name"]

        # Style nodes based on type
        if block["type"] == "function_container":
            mermaid_code += f'    {node_id}["🔧 {node_name}()"]\n'
        elif block["type"] == "class_container":
            mermaid_code += f'    {node_id}["📦 class {node_name}"]\n'
        else:
            mermaid_code += f'    {node_id}["{node_name}"]\n'

        # Add subgraphs for functions and classes if there are multiple
        if len(block["functions"]) > 1 or len(block["classes"]) > 1:
            mermaid_code += f"    subgraph {node_id}_content\n"

            for func in block["functions"]:
                func_id = f"{node_id}_{func}"
                mermaid_code += f'        {func_id}(["function {func}"])\n'
                mermaid_code += f"        {node_id} --> {func_id}\n"

            for cls in block["classes"]:
                cls_id = f"{node_id}_{cls}"
                mermaid_code += f'        {cls_id}(["class {cls}"])\n'
                mermaid_code += f"        {node_id} --> {cls_id}\n"

            mermaid_code += "    end\n"

    # Add connections (sequential flow)
    for i in range(len(analysis) - 1):
        current_id = analysis[i]["id"]
        next_id = analysis[i + 1]["id"]
        mermaid_code += f"    {current_id} --> {next_id}\n"

    return mermaid_code
